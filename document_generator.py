import os
import json
import datetime
import tempfile
from weasyprint import HTML
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import templates
from db import get_db_connection

def validate_parameters(doc_type, parameters):
    """Validate and preprocess parameters before document generation"""
    template_params = get_document_parameters(doc_type)
    validated_params = {}
    
    # Collect all required fields
    required_fields = []
    for section in template_params:
        for field in section['fields']:
            if field.get('required', False):
                required_fields.append(field['id'])
    
    # Check required fields
    missing_fields = [field for field in required_fields if not parameters.get(field)]
    if missing_fields:
        raise ValueError(f"Missing required fields: {', '.join(missing_fields)}")
    
    # Process and sanitize parameters
    for section in template_params:
        for field in section['fields']:
            field_id = field['id']
            if field_id in parameters:
                value = parameters[field_id]
                
                # Type validation and conversion
                if field.get('type') == 'number':
                    try:
                        value = float(value)
                    except (ValueError, TypeError):
                        raise ValueError(f"Field '{field_id}' must be a number")
                elif field.get('type') == 'date':
                    if not isinstance(value, (datetime.date, str)):
                        raise ValueError(f"Field '{field_id}' must be a valid date")
                
                # Options validation
                if field.get('options') and value not in field['options']:
                    raise ValueError(f"Invalid value for field '{field_id}'. Must be one of: {', '.join(field['options'])}")
                
                validated_params[field_id] = value
    
    return validated_params

def get_template_by_type(doc_type):
    """Get the appropriate template for the document type"""
    template_mapping = {
        "nda": templates.nda_template.generate,
        "invoice": templates.invoice_template.generate,
        "letter_of_intent": templates.letter_of_intent_template.generate,
        "proposal": templates.proposal_template.generate,
        "scope_of_work": templates.scope_of_work_template.generate
    }
    
    if doc_type not in template_mapping:
        raise ValueError(f"Document type '{doc_type}' is not supported")
    
    return template_mapping[doc_type]

def get_industry_profile(user_id):
    """Get the user's industry profile"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        cursor.execute("""
            SELECT industry, company_size, business_type, target_market,
                   company_description, document_preferences, brand_colors
            FROM industry_profiles
            WHERE user_id = ?
        """, (user_id,))
        profile = cursor.fetchone()
        
        if not profile:
            return None
            
        return {
            'industry': profile['industry'],
            'company_size': profile['company_size'],
            'business_type': profile['business_type'],
            'target_market': profile['target_market'],
            'company_description': profile['company_description'],
            'document_preferences': json.loads(profile['document_preferences']),
            'brand_colors': json.loads(profile['brand_colors'])
        }
    finally:
        cursor.close()
        conn.close()

def get_industry_specific_content(doc_type, industry):
    """Get industry-specific template modifiers"""
    industry_templates = {
        'Technology & Software': {
            'proposal': {
                'tone': 'technical and innovative',
                'focus_points': ['technological innovation', 'scalability', 'integration capabilities'],
                'terms': ['implementation', 'deployment', 'system integration', 'technical support']
            }
        },
        'Healthcare': {
            'proposal': {
                'tone': 'professional and compliant',
                'focus_points': ['patient care', 'compliance', 'healthcare standards'],
                'terms': ['HIPAA compliance', 'patient data', 'medical protocols']
            }
        },
        'Finance & Banking': {
            'proposal': {
                'tone': 'formal and precise',
                'focus_points': ['security', 'regulatory compliance', 'risk management'],
                'terms': ['ROI', 'financial controls', 'compliance requirements']
            }
        }
        # Add more industries and their specific templates
    }
    
    return industry_templates.get(industry, {}).get(doc_type, {})

def apply_industry_customizations(content, profile, doc_type):
    """Apply industry-specific customizations to document content"""
    if not profile:
        return content
        
    industry_content = get_industry_specific_content(doc_type, profile['industry'])
    
    # Replace generic terms with industry-specific ones
    if 'terms' in industry_content:
        for term in industry_content['terms']:
            if term.lower() not in content.lower():
                content = content.replace(
                    'implementation',
                    f'{term}',
                    1  # Replace only first occurrence to avoid over-modification
                )
    
    # Add industry-specific focus points if not present
    if 'focus_points' in industry_content:
        focus_points = industry_content['focus_points']
        for point in focus_points:
            if point.lower() not in content.lower():
                content += f"\n\nKey {profile['industry']} Focus: {point.capitalize()}"
    
    # Apply brand colors
    if 'brand_colors' in profile:
        content = content.replace(
            '#2E86C1',  # Default primary color in templates
            profile['brand_colors'].get('primary', '#2E86C1')
        ).replace(
            '#2874A6',  # Default secondary color in templates
            profile['brand_colors'].get('secondary', '#2874A6')
        )
    
    return content

def generate_document_content(doc_type, parameters, user_id=None):
    """Generate document content with industry customization"""
    # Validate and preprocess parameters
    validated_params = validate_parameters(doc_type, parameters)
    
    # Get template generator
    template_generator = get_template_by_type(doc_type)
    
    # Generate content with validated parameters
    content = template_generator(validated_params)
    
    # Apply industry customizations if user_id is provided
    if user_id:
        profile = get_industry_profile(user_id)
        if profile:
            content = apply_industry_customizations(content, profile, doc_type)
    
    return content

def generate_pdf(content, title):
    """Generate a PDF file from HTML content"""
    # Create a temporary HTML file
    with tempfile.NamedTemporaryFile(suffix='.html', delete=False) as tmp_html:
        tmp_html.write(content.encode('utf-8'))
        tmp_html_path = tmp_html.name
    
    # Create a temporary PDF file
    pdf_path = tmp_html_path.replace('.html', '.pdf')
    
    # Generate PDF
    HTML(tmp_html_path).write_pdf(pdf_path)
    
    # Clean up the HTML file
    os.unlink(tmp_html_path)
    
    return pdf_path

def generate_docx(content, title, doc_type):
    """Generate a DOCX file from content"""
    # Create a new Document
    doc = Document()
    
    # Add a title
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add the date
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(datetime.datetime.now().strftime("%B %d, %Y"))
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add a line
    doc.add_paragraph("_" * 50)
    
    # This is a simple implementation - in a real app, you'd parse the HTML content
    # and convert it properly to DOCX format
    paragraphs = content.replace('<br>', '\n').replace('<p>', '').replace('</p>', '\n\n')
    paragraphs = paragraphs.replace('<strong>', '').replace('</strong>', '')
    paragraphs = paragraphs.replace('<em>', '').replace('</em>', '')
    paragraphs = paragraphs.replace('&nbsp;', ' ')
    
    # Split by lines and add each as a paragraph
    for para in paragraphs.split('\n'):
        if para.strip():
            doc.add_paragraph(para.strip())
    
    # Save to a temporary file
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
        docx_path = tmp_docx.name
    
    doc.save(docx_path)
    return docx_path

def get_document_parameters(doc_type):
    """Get the parameter definitions for a document type"""
    if doc_type == "nda":
        return templates.nda_template.PARAMETERS
    elif doc_type == "invoice":
        return templates.invoice_template.PARAMETERS
    elif doc_type == "letter_of_intent":
        return templates.letter_of_intent_template.PARAMETERS
    elif doc_type == "proposal":
        return templates.proposal_template.PARAMETERS
    elif doc_type == "scope_of_work":
        return templates.scope_of_work_template.PARAMETERS
    else:
        raise ValueError(f"Document type '{doc_type}' is not supported")
